from __future__ import annotations

from pathlib import Path


def parse_docx(path: Path | None = None, data: bytes | None = None) -> str:
    """从 .docx 抽取段落文本。依赖 python-docx。"""
    try:
        from docx import Document
    except ImportError:
        return ""
    if path is not None:
        doc = Document(str(path))
    elif data is not None:
        from io import BytesIO

        doc = Document(BytesIO(data))
    else:
        return ""
    parts = [p.text for p in doc.paragraphs if p.text and p.text.strip()]
    return "\n".join(parts).strip()


def parse_file(path: Path) -> str:
    return parse_docx(path=path)


__all__ = ["parse_docx", "parse_file"]
